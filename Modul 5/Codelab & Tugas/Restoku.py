import os
from docx import Document

# File untuk menyimpan data pesanan
DATA_FILE = "pesanan.txt"
DOCX_FILE = "pesanan.docx"

# Menu makanan dan minuman beserta path gambar
MAKANAN = [
    {"name": "Ayam Krispi", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Ayam Krispi.jpg"},
    {"name": "Ayam Kremes", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Ayam Kremes.jpg"},
    {"name": "Ayam Geprek", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Ayam Geprek.png"},
    {"name": "Soto Ayam", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Soto Ayam.jpg"},
    {"name": "Lalapan Ayam", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Lalapan Ayam.jpg"},
    {"name": "Mie Ayam", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Mie Ayam.jpg"},
    {"name": "Sate Ayam", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Sate Ayam.jpg"},
    {"name": "Rendang Sapi", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Rendang Sapi.jpg"},
    {"name": "Nasi Goreng", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Nasi Goreng.jpg"},
    {"name": "Bakso", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Makanan\Bakso.jpg"},
]

MINUMAN = [
    {"name": "Es Cendol", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Minuman\Es Cendol.jpg"},
    {"name": "Es Teh", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Minuman\Es Teh.jpg"},
    {"name": "Thai Tea", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Minuman\Thai Tea.jpg"},
    {"name": "Matcha Latte", "image": r"D:\Kuliah\3. Semester 3\Proglan\Praktikum\Modul 5\Tugas\Minuman\Mathca Latte.jpg"},
]

# Fungsi untuk membaca pesanan dari file
def read_orders():
    if not os.path.exists(DATA_FILE):
        return []
    with open(DATA_FILE, "r") as file:
        lines = file.readlines()
        orders = [eval(line.strip()) for line in lines]
    return orders

# Fungsi untuk menyimpan pesanan ke file
def save_orders(orders):
    with open(DATA_FILE, "w") as file:
        for order in orders:
            file.write(f"{order}\n")

# Fungsi untuk memperbarui file .docx setelah setiap perubahan
def update_docx(orders):
    doc = Document()
    doc.add_heading("Data Pesanan Restoran", level=1)
    for i, order in enumerate(orders):
        doc.add_paragraph(f"{i + 1}. Nama Pelanggan: {order['name']}")
        doc.add_paragraph("   Pesanan:")
        for item in order["items"]:
            doc.add_paragraph(f"     - {item['name']} (Gambar: {item['image']})")
        doc.add_paragraph(f"   Status: {order['status']}")
        doc.add_paragraph("")
    doc.save(DOCX_FILE)
    print(f"Pesanan berhasil disimpan ke '{DOCX_FILE}'.")

# Fungsi untuk menampilkan menu makanan dan minuman
def display_menu():
    print("\n=== Menu Makanan ===")
    for i, item in enumerate(MAKANAN, 1):
        print(f"{i}. {item['name']}")
    print("\n=== Menu Minuman ===")
    for i, item in enumerate(MINUMAN, 1):
        print(f"{i}. {item['name']}")

# Fungsi untuk membuat pesanan baru
def create_order():
    print("\n=== Buat Pesanan Baru ===")
    name = input("Nama Pelanggan: ")
    display_menu()
    food_choices = input("Pilih nomor makanan (pisahkan dengan koma): ").split(",")
    drink_choices = input("Pilih nomor minuman (pisahkan dengan koma): ").split(",")

    try:
        items = [{"name": MAKANAN[int(i.strip()) - 1]["name"], 
                  "image": MAKANAN[int(i.strip()) - 1]["image"]} for i in food_choices if i.strip().isdigit()] + \
                [{"name": MINUMAN[int(i.strip()) - 1]["name"], 
                  "image": MINUMAN[int(i.strip()) - 1]["image"]} for i in drink_choices if i.strip().isdigit()]
    except IndexError:
        print("Error: Pilihan tidak valid.")
        return

    new_order = {
        "name": name,
        "items": items,
        "status": "diproses"
    }
    orders = read_orders()
    orders.append(new_order)
    save_orders(orders)
    update_docx(orders)
    print("Pesanan berhasil dibuat dan disimpan ke file!")

# Menu utama
def main():
    while True:
        print("\n=== Sistem Manajemen Pesanan Restoran ===")
        print("1. Buat Pesanan")
        print("2. Tampilkan Pesanan")
        print("3. Keluar")
        choice = input("Pilih menu: ")
        if choice == "1":
            create_order()
        elif choice == "2":
            orders = read_orders()
            if not orders:
                print("Tidak ada pesanan.")
            else:
                for i, order in enumerate(orders):
                    print(f"{i + 1}. {order}")
        elif choice == "3":
            print("Terima kasih telah menggunakan sistem!")
            break
        else:
            print("Pilihan tidak valid.")

if __name__ == "__main__":
    main()
