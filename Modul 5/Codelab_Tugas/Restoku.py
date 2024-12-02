import os
from docx import Document
from docx.shared import Inches

DATA_FILE = "pesanan.txt"
DOCX_FILE = "pesanan.docx"

MAKANAN = [
    "Ayam Krispi", "Ayam Kremes", "Ayam Geprek", "Soto Ayam",
    "Lalapan Ayam", "Mie Ayam", "Sate Ayam", "Rendang Sapi",
    "Nasi Goreng", "Bakso"
]

MINUMAN = ["Es Cendol", "Es Teh", "Thai Tea", "Matcha Latte"]

def read_orders():
    if not os.path.exists(DATA_FILE):
        return []
    with open(DATA_FILE, "r") as file:
        lines = file.readlines()
        return [eval(line.strip()) for line in lines]

def save_orders(orders):
    with open(DATA_FILE, "w") as file:
        for order in orders:
            file.write(f"{order}\n")

def update_docx(orders):
    doc = Document()
    doc.add_heading("Data Pesanan Restoran", level=1)
    for i, order in enumerate(orders):
        doc.add_paragraph(f"{i + 1}. Nama Pelanggan: {order['name']}")
        doc.add_paragraph("   Pesanan:")
        for item in order["items"]:
            doc.add_paragraph(f"     - {item['name']}")
            
            if os.path.exists(item["image"]):
                doc.add_picture(item["image"], width=Inches(2))
            else:
                doc.add_paragraph("       (Gambar tidak ditemukan)")
        doc.add_paragraph(f"   Status: {order['status']}")
        doc.add_paragraph("")
    doc.save(DOCX_FILE)
    print(f"Pesanan berhasil disimpan ke '{DOCX_FILE}'.")

def display_menu():
    print("\n=== Menu Makanan ===")
    for i, item in enumerate(MAKANAN, 1):
        print(f"{i}. {item}")
    print("\n=== Menu Minuman ===")
    for i, item in enumerate(MINUMAN, 1):
        print(f"{i}. {item}")


def create_order():
    print("\n=== Buat Pesanan Baru ===")
    name = input("Nama Pelanggan: ")
    display_menu()
    food_choices = input("Pilih nomor makanan : ").split(",")
    drink_choices = input("Pilih nomor minuman : ").split(",")
    items = []
    try:
        items += [{"name": MAKANAN[int(i.strip()) - 1], "image": input(f"Path gambar untuk {MAKANAN[int(i.strip()) - 1]}: ")} for i in food_choices if i.strip().isdigit()]
        items += [{"name": MINUMAN[int(i.strip()) - 1], "image": input(f"Path gambar untuk {MINUMAN[int(i.strip()) - 1]}: ")} for i in drink_choices if i.strip().isdigit()]
    except (IndexError, ValueError):
        print("Input tidak valid!")
        return

    new_order = {"name": name, "items": items, "status": "diproses"}
    orders = read_orders()
    orders.append(new_order)
    save_orders(orders)
    update_docx(orders)
    print("Pesanan berhasil dibuat!")

def read_all_orders():
    orders = read_orders()
    if not orders:
        print("Tidak ada pesanan.")
    else:
        for i, order in enumerate(orders, 1):
            print(f"{i}. {order}")

def update_order():
    orders = read_orders()
    if not orders:
        print("Tidak ada pesanan.")
        return
    read_all_orders()
    try:
        index = int(input("Pilih nomor pesanan yang ingin diubah: ")) - 1
        if 0 <= index < len(orders):
            status = input("Ubah status menjadi (diproses/selesai/batal): ").strip().lower()
            if status in ["diproses", "selesai", "batal"]:
                orders[index]["status"] = status
                save_orders(orders)
                update_docx(orders)
                print("Status pesanan berhasil diubah!")
            else:
                print("Status tidak valid.")
        else:
            print("Nomor pesanan tidak valid.")
    except ValueError:
        print("Input tidak valid.")

def delete_order():
    orders = read_orders()
    if not orders:
        print("Tidak ada pesanan.")
        return
    read_all_orders()
    try:
        index = int(input("Pilih nomor pesanan yang ingin dihapus: ")) - 1
        if 0 <= index < len(orders):
            if orders[index]["status"] == "batal":
                orders.pop(index)
                save_orders(orders)
                update_docx(orders)
                print("Pesanan berhasil dihapus!")
            else:
                print("Hanya pesanan dengan status 'batal' yang bisa dihapus.")
        else:
            print("Nomor pesanan tidak valid.")
    except ValueError:
        print("Input tidak valid.")


def search_order():
    orders = read_orders()
    if not orders:
        print("Tidak ada pesanan.")
        return
    name = input("Masukkan nama pelanggan yang ingin dicari: ").strip()
    results = [order for order in orders if order["name"].lower() == name.lower()]
    if results:
        for order in results:
            print(order)
    else:
        print("Pesanan tidak ditemukan.")

def main():
    while True:
        print("\n=== Sistem Manajemen Pesanan Restoran ===")
        print("1. Buat Pesanan")
        print("2. Tampilkan Semua Pesanan")
        print("3. Ubah Status Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan")
        print("6. Keluar")
        choice = input("Pilih menu: ")
        if choice == "1":
            create_order()
        elif choice == "2":
            read_all_orders()
        elif choice == "3":
            update_order()
        elif choice == "4":
            delete_order()
        elif choice == "5":
            search_order()
        elif choice == "6":
            print("Terima kasih telah menggunakan sistem!")
            break
        else:
            print("Pilihan tidak valid.")

if __name__ == "__main__":
    main()
